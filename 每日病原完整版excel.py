import os
from docx import Document
from openpyxl import Workbook
import numpy as np
import pandas as pd
import random
import datetime
import time
# 获取文件夹
path=r'D:\Users\Administrator\Desktop\4月病原'
files=os.listdir(path)
allfile=[]
for i in files:
    allfile.append(os.path.join(path,i))
# print(allfile)
zongbiao=[]
# 遍历文件
for i in allfile:

    日期=str(i).split('\\')[-1]
    biao=Document(i)

    # print(biao.tables[-1].rows)
    table=[]
    for row in biao.tables[-1].rows:
        line=[]
        for i in row.cells:
            line.append(i.text)
        # print(line)
        table.append(line)
    # print(table)
    a=np.array(table[2:])
    # 表头1=['序号','样本信息','fam','fam值','hex','hex值','备注']
    # 表头2=['序号','样本信息','fam','fam值','备注']
    s=pd.DataFrame(a)
    s.drop(s.columns[4:len(s.columns)],axis=1,inplace=True)
    year=日期[:8][:4]
    month=日期[:8][4:6]
    day=日期[:8][6:]
    shijian=year+'-'+month+'-'+day
    da=datetime.datetime.strptime(shijian,"%Y-%m-%d")
    # print(da)
    s['日期']=da
    if str(日期[8:]).split('检')[0]=='非瘟':
        检测类型='非洲猪瘟'
    elif '冠状' in str(日期[8:]).split('检')[0]:
        检测类型 = '猪δ冠状病毒'
    elif '蓝耳' in str(日期[8:]).split('检')[0]:
        检测类型 = '蓝耳'
    elif '圆环' in str(日期[8:]).split('检')[0]:
        检测类型 = '圆环'
    elif '细小' in str(日期[8:]).split('检')[0]:
        检测类型 = '细小'
    elif '伪狂' in str(日期[8:]).split('检')[0]:
        检测类型 = '伪狂犬'
    elif '腹泻' in str(日期[8:]).split('检')[0]:
        if '三重腹泻' in str(日期[8:]).split('检')[0]:
            检测类型 = '三重腹泻'
        else:
            检测类型 ='腹泻'
    elif '猪瘟' in str(日期[8:]).split('检')[0]:
        检测类型 = '猪瘟'
    else:
        检测类型= str(日期[8:]).split('检')[0]
    s['检测类型']=检测类型
    s.drop(index=[len(s)-1],inplace=True)
    s.drop(index=[len(s)-1],inplace=True)
    # print(s)
    zongbiao.append(s)
    # print(s)
print("休息一会。。。。。")
allbiao=pd.concat(zongbiao)
# print(allbiao)

col_name=['序号','检测样本信息','结果','FAM值','日期','检测类型']
allbiao.columns=col_name
# allbiao.to_excel('b.xlsx')

path_name=r'D:\Users\Administrator\Desktop\模板1\联系人.xlsx'



场名=pd.read_excel(path_name,sheet_name='Sheet2')
# 获取检测类型




all_samp=allbiao
部门=场名['部门']
场区1=场名['场区']
公司1=场名['公司']

测试=all_samp['检测样本信息']
日期=all_samp['日期']
序号1=all_samp['序号']
结果1=all_samp['结果']
FAM值1=all_samp['FAM值']
检测类型=all_samp['检测类型']



d_部门=[]
d_样品=[]
d_日期=[]
d_序号=[]
d_结果=[]
d_FAM值=[]
d_检测类型=[]


for i in 部门:
    d_部门.append(i)
for i in 测试:
    d_样品.append(i)
for i in 日期:
    d_日期.append(i)
for i in 序号1:
    d_序号.append(i)
for i in 结果1:
    d_结果.append(i)
for i in FAM值1:
    d_FAM值.append(i)
for i in 检测类型:
    d_检测类型.append(i)



# print(样品)
场区=[] #样本对应的场区
样本=[] #有部门的样品
序号=[] #有部门样品的index
样品类型=[]
检测类型=[]
结果=[]
fam值=[]

day=[]
试剂盒=[]
公司=[]
# no_部门=[]
# no_序号=[]
# 填写部门

def kongchangqu(x):
    if x in i:
        场区.append(x)
        公司.append(公司1[d_部门.index(x)])
        样本.append(i)
        序号.append(d_序号[d_样品.index(i)])
        day.append(d_日期[d_样品.index(i)])

        检测类型.append(d_检测类型[d_样品.index(i)])
        结果.append(d_结果[d_样品.index(i)])
        fam值.append(d_FAM值[d_样品.index(i)])

for i in d_样品:
    if '工程部' in i and '阜阳' not in i:
        场区.append('工程部')
        公司.append('生物安全工程部')
        样本.append(i)
        序号.append(d_序号[d_样品.index(i)])
        day.append(d_日期[d_样品.index(i)])
        检测类型.append(d_检测类型[d_样品.index(i)])
        结果.append(d_结果[d_样品.index(i)])
        fam值.append(d_FAM值[d_样品.index(i)])
    else:
        for j in d_部门:
            kongchangqu(j)



for i in 样本:
    if '人' in i:
        样品类型.append('人员样本')
   # '血'or'鼻'or'肛'or'精'or'咽'or'粪'or'口'or'肉'or'死猪'
    elif '血' in i or '鼻' in i or '肛'in i or '精'in i or '咽' in i or '粪'in i or '口'in i or '肉' in i or '死猪' in i or' 唾' in i :
        样品类型.append('猪源样本')
    else:
        样品类型.append('环境样本')

for i in all_samp['检测类型']:
    if '非瘟' in i or '非洲猪瘟' in i:
        试剂盒.append('青岛立见')
    elif '三重腹泻' in i:
        试剂盒.append('世纪元亨')
    else:
        试剂盒.append('维特康')



dir1={'id':序号,
      '日期':day,
      '公司':公司,
      '部门':场区,
      '样本类型':样品类型,
      '检测类型':检测类型,
      '试剂盒':试剂盒,
      '检测样本信息':样本,
      '结果':结果,
      'FAM值':fam值,
      }
df=pd.DataFrame(dir1)
# df.to_excel('a.xlsx')
aa = time.strftime("%Y-%m-%d", time.localtime())
name = str(aa)+'.xlsx'
path1 = r'D:\Users\Administrator\Desktop'
fil = os.path.join(path1, name)
总计1=len(样本)
总计2 = len(allbiao['检测样本信息'])
if 总计1==总计2:
    print('非常好，一切顺利！！')
    df.to_excel(fil)
else:
    print('还不行，需要改！！')
    # df.to_excel(fil)
    allbiao.to_excel('未完成.xlsx')
print('over!')
# def kongchangqu(x):
#     if x in  样品[i]:
#         场区.append(x)
#         公司.append(公司1[d_部门.index(x)])
#         样本.append(样品[i])
#         序号.append(i)
#         no_部门.remove(样品[i])
#         no_序号.remove(i)
#
# konglist=['巩店公猪站','马店母猪场','李灿']
# 部门1=d_部门[:-1]
# for i in range(len(样品)):
#     if i not in 序号:
#         no_部门.append(样品[i])
#         no_序号.append(i)
#         for x in 部门1:
#             kongchangqu(x)
#
# a=10
# while a>0:
#     a-=1
#     for i in no_序号:
#         if '工程部' in 样品[i] or '生物安全' in 样品[i]:
#             场区.append('工程部')
#             样本.append(样品[i])
#             公司.append('生物安全工程部')
#             序号.append(i)
#             del no_部门[(no_序号.index(i))]
#             no_序号.remove(i)
#
# print(no_部门)
# print(no_序号)
# for i in range(len(样品)):
#     if i not in no_序号:
#         no_部门.append(样品[i])
#         no_序号.append(i)
#         if '工程' in 样品[i]:
#             场区.append('工程部')
#             公司.append(公司1[d_部门.index(i)])
#             样本.append(样品[i])
#             序号.append(i)
#             no_部门.remove(样品[i])
#             no_序号.remove(i)


# for i in 样品:
#     if '人' in i:
#         样品类型.append('人员样本')
#    # '血'or'鼻'or'肛'or'精'or'咽'or'粪'or'口'or'肉'or'死猪'
#     elif '血' in i or '鼻' in i or '肛'in i or '精'in i or '咽' in i or '粪'in i or '口'in i or '肉' in i or '死猪' in i or' 唾' in i :
#         样品类型.append('猪源样本')
#     else:
#         样品类型.append('环境样本')
#
# # print(样品类型)
# # 填写试剂盒
# for i in all_samp['检测类型']:
#     if '非瘟' in i or '非洲猪瘟' in i:
#         试剂盒.append('青岛立见')
#     elif '三重腹泻' in i:
#         试剂盒.append('世纪元亨')
#     else:
#         试剂盒.append('维特康')
# # 填写场区




#判断是否有样品没有录进去，如果都录进去了就成功，没录进去就检查一下
# if len(no_部门)==0:
#     dir = {'序号': 序号, '样品': 样本, '场区': 场区,'公司':公司}
#     df = pd.DataFrame(dir)
#     df = df.drop_duplicates(subset=['序号'], keep='first')
#     df = df.dropna(axis=0, how='all', thresh=None, subset=None, inplace=False)  # 删除全部为空的行
#     part = []
#     for i in df['场区']:
#         part.append(i)
#     gongsi=[x for x in df['公司']]
#     allbiao['部门'] = part
#     allbiao['样本类型']=样品类型
#     allbiao['试剂盒']=试剂盒
#     allbiao['公司']=gongsi
#     allbiao['id']=''
#     allbiao=allbiao[['id','日期','公司','部门','样本类型','检测类型','试剂盒','检测样本信息','结果','FAM值']]
#     # df=pd.concat([allbiao,df],axis=1)
#     # df.to_excel('完成.xlsx')
#
#     aa = time.strftime("%Y-%m-%d", time.localtime())
#     name = str(aa)+'.xlsx'
#     path1 = r'D:\Users\Administrator\Desktop'
#     fil = os.path.join(path1, name)
#     总计2 = len(allbiao['检测样本信息'])
#     if 总计1==总计2:
#         print('非常好，一切顺利！！')
#     else:
#         print('还不行，需要改！！')
#     allbiao.to_excel(fil)
#     print('over!')
# else:
#     no_部门={'样品':no_部门,'序号':no_序号}
#     df_no=pd.DataFrame(no_部门)
#     df_no.to_excel('没录入的样本.xlsx')
#     dir = {'序号': 序号, '样品': 样本, '场区': 场区}
#     df = pd.DataFrame(dir)
#     df = df.drop_duplicates(subset=['序号'], keep='first')
#     df.to_excel('没完成1.xlsx')
#
#     # df=pd.concat([df,allbiao],axis=1,join_axes=[df.index])
#     allbiao.to_excel('没完成.xlsx')




