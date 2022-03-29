import os
from docx import Document
from openpyxl import Workbook
import numpy as np
import pandas as pd
import random
import datetime
path=r'D:\Users\Administrator\Desktop\3月病原'
files=os.listdir(path)
allfile=[]
for i in files:
    allfile.append(os.path.join(path,i))
# print(allfile)
zongbiao=[]
for i in allfile:

    日期=str(i).split('\\')[-1]
    biao=Document(i)

    # print(biao.tables[-1].rows)
    table=[]
    for row in biao.tables[-1].rows:
        line=[]
        for i in row.cells:
            line.append(i.text)
        # print(line)
        table.append(line)
    # print(table)
    a=np.array(table[2:])
    # 表头1=['序号','样本信息','fam','fam值','hex','hex值','备注']
    # 表头2=['序号','样本信息','fam','fam值','备注']
    s=pd.DataFrame(a)
    s.drop(s.columns[4:len(s.columns)],axis=1,inplace=True)
    year=日期[:8][:4]
    month=日期[:8][4:6]
    day=日期[:8][6:]
    shijian=year+'-'+month+'-'+day
    da=datetime.datetime.strptime(shijian,"%Y-%m-%d")
    # print(da)
    s['日期']=da
    if str(日期[8:]).split('检')[0]=='非瘟':
        检测类型='非洲猪瘟'
    elif '冠状' in str(日期[8:]).split('检')[0]:
        检测类型 = '猪δ冠状病毒'
    elif '蓝耳' in str(日期[8:]).split('检')[0]:
        检测类型 = '蓝耳'
    else:
        检测类型= str(日期[8:]).split('检')[0]
    s['检测类型']=检测类型
    s.drop(index=[len(s)-1],inplace=True)
    s.drop(index=[len(s)-1],inplace=True)
    # print(s)
    zongbiao.append(s)
    # print(s)
print("休息一会。。。。。")
allbiao=pd.concat(zongbiao)
name='总计.xlsx'
path1=r'D:\Users\Administrator\Desktop'
fil=os.path.join(path1,name)

allbiao.to_excel(fil)
print('over!')
